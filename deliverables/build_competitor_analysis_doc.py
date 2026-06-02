from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_SCRIPTS = Path(
    "/Users/seyiadelaju/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.430.10722/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights

OUT = Path("/Users/seyiadelaju/Projects/trustgrid/deliverables/trustgrid-competitor-analysis.docx")

NAVY = RGBColor(15, 23, 42)
BLUE = RGBColor(30, 64, 175)
MUTED = RGBColor(71, 85, 105)
LIGHT_BLUE = "EEF2FF"
LIGHT_GRAY = "F1F5F9"
AMBER = "FEF3C7"
ROSE = "FFE4E6"
GREEN = "DCFCE7"


def style_run(run, *, size=10, bold=False, italic=False, color=NAVY):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text, *, bold=False, color=NAVY, size=8.6):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    style_run(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, weights, *, header_fill=LIGHT_BLUE, size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=size)
        set_cell_fill(table.rows[0].cells[idx], header_fill)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=size)
    widths = column_widths_from_weights(weights)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    return table


def add_note(doc, title, text, *, fill=AMBER):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{title}: ")
    style_run(run, size=9.4, bold=True)
    run = paragraph.add_run(text)
    style_run(run, size=9.4)
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=0)
    return table


def add_heading(doc, text, level=1, *, page_break=False):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def add_body(doc, text, *, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        style_run(run, size=9.6, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        style_run(run, size=9.6)
    else:
        style_run(paragraph.add_run(text), size=9.6)
    return paragraph


def add_bullet(doc, text, *, bold_prefix=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.22)
    if bold_prefix and text.startswith(bold_prefix):
        style_run(paragraph.add_run(bold_prefix), size=9.5, bold=True)
        style_run(paragraph.add_run(text[len(bold_prefix):]), size=9.5)
    else:
        style_run(paragraph.add_run(text), size=9.5)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.22)
    style_run(paragraph.add_run(text), size=9.5)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(paragraph.add_run("TrustGrid competitor analysis  |  "), size=8, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(30, 41, 59)
    styles["Subtitle"].font.name = "Arial"
    styles["Subtitle"].font.size = Pt(11)
    styles["Subtitle"].font.color.rgb = MUTED
    for name, size in (("Heading 1", 15), ("Heading 2", 11.5)):
        styles[name].font.name = "Arial"
        styles[name].font.bold = True
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = BLUE

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("TRUSTGRID  |  COMPETITIVE REALITY CHECK"), size=8, bold=True, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    doc.add_heading("TrustGrid Competitor Analysis", 0)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A skeptical market assessment and recommended wedge")
    add_table(
        doc,
        ["Prepared for", "Research date", "Purpose"],
        [["TrustGrid founding team", "2 June 2026", "Product and go-to-market decisions"]],
        [1.6, 1.2, 3.7],
        size=9,
    )
    doc.add_paragraph()
    add_note(
        doc,
        "Method",
        "This memo uses TrustGrid's internal product documents and public information from competitor websites. "
        "Competitor claims are treated as vendor claims, not independently audited facts. The aim is to identify what must be proven, not to decorate a pitch.",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "Executive Summary", 1)
    add_note(
        doc,
        "Bottom line",
        "TrustGrid is not entering an empty market. The current vision overlaps with local estate-management products, "
        "global volunteer-management platforms, contractor-compliance software, and identity-verification APIs. "
        "The opportunity may still be real, but the company does not yet have a moat. It has a product hypothesis and a potentially valuable distribution wedge.",
        fill=AMBER,
    )
    add_bullet(doc, "The strongest initial wedge is not a broad smart-city operating system. It is trust-gated membership, credentialing, and deployment for one community network, starting with a specific RCCG province or convention workflow.")
    add_bullet(doc, "Estate management is crowded. Gate Africa and Sylogate already sell access control and community operations locally. MyGate shows how broad and sophisticated that category can become.")
    add_bullet(doc, "Volunteer coordination and contractor compliance are mature categories globally. Rosterfy and ComplyFlow already cover substantial parts of the proposed roadmap.")
    add_bullet(doc, "Identity verification should mostly be bought, not built. Prembly and Youverify compress a feature that earlier TrustGrid documents sometimes imply is differentiating.")
    add_bullet(doc, "The defensible hypothesis is a portable Trust Passport backed by community hierarchy and authority chains. That is not a moat until real institutions issue credentials, workers reuse them, and another institution changes a decision because of them.")

    add_heading(doc, "1. Product Definition Used for This Analysis", 1)
    add_body(doc, "TrustGrid's canonical June 2026 positioning is Community Trust Infrastructure: a platform for communities to establish, verify, govern, and reuse trust signals. The proposed modules include community hierarchy, delegated authority, memberships, endorsements, portable Trust Passports, credentials, opportunities, and QR verification.")
    add_body(doc, "The current MVP is narrower: worker registry, simulated identity verification, trust scoring, endorsements, service requests, assignments, reviews, incidents, and an institution dashboard. Several vision-critical capabilities are not yet proven in production.")
    add_table(
        doc,
        ["Layer", "Current reality", "What remains unproven"],
        [
            ["MVP workflow", "Registry, assignments, reviews, incidents, basic score loop", "Repeated real-world usage and operational reliability"],
            ["Verification", "Simulated identity adapter in MVP", "Live provider integration, cost, consent, false-positive handling"],
            ["Community graph", "Documented architecture and positioning", "Real hierarchy onboarding and delegated-authority workflows"],
            ["Trust Passport", "Core product concept", "Cross-institution reuse that changes an access, staffing, or hiring decision"],
            ["Distribution", "RCCG proposed as beachhead", "Named buyer, budget owner, procurement path, and rollout commitment"],
        ],
        [1.3, 2.7, 3.3],
    )

    add_heading(doc, "2. The Five Most Relevant Platforms", 1, page_break=True)
    add_body(doc, "These are not all identical competitors. They are the products a skeptical buyer or investor will use to test whether TrustGrid is necessary.")
    add_table(
        doc,
        ["Platform", "Category", "Overlap with TrustGrid", "Why it matters"],
        [
            ["Gate Africa", "Nigeria estate operations", "Estate access, resident workflows, event passes, payments", "A local buyer may prefer one estate operating app instead of adding another platform."],
            ["Sylogate", "Nigeria estate security", "Residents, staff, visitors, QR access, incidents, multi-estate access", "Direct evidence that local teams already sell community trust and access features."],
            ["MyGate", "Large-scale community ERP", "Recurring staff, vendors, contractors, approvals, audit logs, payroll-oriented staff records", "Shows the estate category can expand far beyond visitor codes and become hard to displace."],
            ["Rosterfy", "Volunteer management", "Recruitment, screening, onboarding, credentials, scheduling, attendance, reporting", "Strong benchmark for church conventions and volunteer deployments."],
            ["ComplyFlow", "Contractor compliance", "Prequalification, worker records, credentials, incidents, permits, site access, sign-in", "Strong benchmark for facilities, schools, vendors, and multi-site compliance."],
        ],
        [1.05, 1.35, 2.35, 2.55],
        size=8.4,
    )

    add_heading(doc, "Direct Competitors: Local Estate Software", 2)
    add_body(doc, "Gate Africa positions itself as an all-in-one estate platform for property managers, owners, and tenants across Africa. Its official site highlights access codes, event-specific invites, community features, payments, support for security teams, and claimed traction of more than one million invites across more than six cities.")
    add_body(doc, "Sylogate is a Lagos-based estate access and management platform. Its official site describes QR passes, staff and visitor management, incident reporting, role-based access, multi-estate switching, and outage fallback codes. It claims more than 20 communities, more than 5,000 residents, and more than 500,000 approved visitors.")
    add_note(
        doc,
        "Implication",
        "Do not pitch TrustGrid as an estate-management app unless the product is intentionally narrowed to win that market. "
        "If estates remain a target, TrustGrid needs a crisp integration story: it should supply trusted credentials and portable service-provider history to estate systems, not rebuild dues, visitor invites, and resident communications.",
        fill=ROSE,
    )

    add_heading(doc, "Category Benchmark: MyGate", 2)
    add_body(doc, "MyGate's official community and visitor-management pages describe resident approvals, guard validation, recurring domestic staff, vendor documentation, contractor passes, staff rosters, access rules, audit logs, and ERP features. It says more than 27,000 communities use the product and reports 1.2 billion entries verified yearly. Treat those as company-reported metrics, but take the product breadth seriously.")
    add_body(doc, "MyGate is not a near-term Nigerian go-to-market threat in every segment. It is a strategic warning: access control, domestic-help records, vendors, payments, and community administration naturally converge.")

    add_heading(doc, "Vertical Benchmark: Rosterfy", 2)
    add_body(doc, "Rosterfy's official site describes volunteer recruitment, onboarding, compliance checks, training, scheduling, attendance, communications, rewards, and reporting. It serves nonprofits, cities, universities, hospitals, sports organizations, and major events. Its onboarding page explicitly covers role-specific workflows, screening, background checks, and expiring credentials.")
    add_body(doc, "For an RCCG convention use case, Rosterfy is much closer than a generic HR tool. TrustGrid must explain why community hierarchy, authority-backed endorsements, and portable credentials matter enough to beat a mature volunteer product or justify a local alternative.")

    add_heading(doc, "Vertical Benchmark: ComplyFlow", 2)
    add_body(doc, "ComplyFlow's official site presents a broad contractor-management suite: supplier prequalification, workforce compliance, incident management, audits, permits, site documents, site access, and visitor sign-in. Its workforce page emphasizes role-, contractor-, and site-specific requirements, live site status, sign-in records, qualifications, and competencies.")
    add_body(doc, "For facilities, schools, and vendor governance, TrustGrid should assume buyers can purchase compliance software. A generic vendor registry is not differentiation.")

    add_heading(doc, "3. Capability Reality Check", 1, page_break=True)
    add_table(
        doc,
        ["Capability", "Existing evidence in market", "TrustGrid assessment"],
        [
            ["Visitor and gate access", "Gate Africa, Sylogate, MyGate", "Commodity. Integrate or deprioritize."],
            ["Resident and estate operations", "Gate Africa, Sylogate, MyGate", "Crowded adjacency. Avoid broad estate ERP scope."],
            ["Volunteer onboarding and scheduling", "Rosterfy", "Mature category. Compete only through local workflow and distribution advantage."],
            ["Contractor and vendor compliance", "ComplyFlow, Appruv", "Mature category. Basic records and document uploads are table stakes."],
            ["Identity and background checks", "Prembly, Youverify, Checkr", "Buy through adapters. Verification alone is not a moat."],
            ["Incident records and performance history", "Present in several categories", "Useful, but insufficiently differentiated by itself."],
            ["Portable Trust Passport", "Not established by reviewed direct competitors in TrustGrid's intended form", "Potential wedge, but only if institutions recognize and reuse it."],
            ["Community authority chain", "Partially adjacent to membership and identity products", "Potentially distinctive in faith and association networks; needs field validation."],
            ["Cross-community trust graph", "A hypothesis, not a proven product advantage", "Potential moat later. Currently an aspiration."],
        ],
        [1.6, 2.55, 2.75],
        size=8.35,
    )

    add_heading(doc, "4. Non-Software Competitors", 1)
    add_body(doc, "The hardest competitor is not another startup. It is the current operating model.")
    add_table(
        doc,
        ["Alternative", "Why customers keep using it", "Why switching is difficult"],
        [
            ["WhatsApp groups", "Free, familiar, already installed, flexible", "TrustGrid adds process and data-entry work before the value is obvious."],
            ["Spreadsheets", "Cheap, locally controlled, easy to customize", "An operator can tolerate messy records if incidents are rare or management turnover is low."],
            ["Personal referrals", "Fast and socially legible", "A score may feel less trustworthy than a known person's recommendation."],
            ["Staffing agencies", "They absorb sourcing and coordination work", "Institutions may pay a premium to avoid operating the system themselves."],
            ["Existing church hierarchy", "Authority and trust already exist offline", "Software must strengthen the hierarchy without creating political friction."],
        ],
        [1.3, 2.7, 2.9],
    )
    add_note(
        doc,
        "YC-style test",
        "If a pilot user does not feel acute pain when the product is removed for two weeks, the product has not found a must-have workflow.",
        fill=AMBER,
    )

    add_heading(doc, "5. Claims to Stop Making", 1, page_break=True)
    add_body(doc, "Several internal documents contain claims that are too confident for the current evidence. They may become true, but they should be treated as hypotheses until measured.")
    add_table(
        doc,
        ["Claim", "Why it is weak today", "Replace with"],
        [
            ["'The market is largely unserved.'", "Estate, volunteer, compliance, and verification products already cover major parts of the surface area.", "'The workflow is fragmented, and we are testing whether community authority can unify it.'"],
            ["'No competitor can replicate RCCG distribution.'", "Access is not distribution until a decision-maker authorizes a rollout and users adopt the product.", "'RCCG may offer a low-cost distribution wedge if we secure a province-level pilot.'"],
            ["'Trust data is a moat.'", "There is no moat before meaningful usage, reuse, consent, and data quality.", "'Repeated credential reuse could create defensibility if institutions rely on the history.'"],
            ["'Algorithmic trust score.'", "Opaque scores create fairness, dispute, and reputational risks. More math does not automatically create trust.", "'Explainable trust evidence with clear appeals and institution-controlled policy.'"],
            ["'Smart-city operating system.'", "The phrase is too broad before one narrow workflow works repeatedly.", "'Trust-gated community membership and deployment, starting with one network.'"],
            ["'10x cheaper' or large ROI estimates", "The internal assumptions are not backed by measured customer baselines.", "'Pilot measurement will compare admin hours, repeat staffing rate, incident handling time, and verification cost.'"],
        ],
        [1.55, 2.8, 2.55],
        size=8.3,
    )

    add_heading(doc, "6. Where TrustGrid Could Still Win", 1)
    add_body(doc, "There is a credible company here if the team is disciplined about the wedge.")
    add_bullet(doc, "Community hierarchy as product infrastructure. RCCG has real parent-child structures, delegated authority, and membership verification needs that generic workforce tools may model awkwardly.")
    add_bullet(doc, "Portable credentials within a bounded network. A Trust Passport can be valuable if a member receives a credential in one parish or event and another authorized node accepts it later.")
    add_bullet(doc, "Opportunity distribution through existing trust relationships. The product may reduce the cost of mobilizing known, eligible people without becoming an open marketplace.")
    add_bullet(doc, "Local operating details. WhatsApp-assisted onboarding, low-bandwidth flows, identity-provider adapters, dispute handling, and field support can matter more than feature count.")
    add_bullet(doc, "Institutional network expansion. A province-level deployment can be a stronger starting unit than selling isolated small organizations one at a time.")
    add_note(
        doc,
        "Narrow positioning recommendation",
        "TrustGrid is a trust-gated membership, credentialing, and deployment network for organized communities. "
        "Start with one RCCG province or convention operation. Prove that credentials are issued, reused, and operationally valuable before expanding into estates, schools, procurement, or smart-city language.",
        fill=GREEN,
    )

    add_heading(doc, "7. Recommended 90-Day Proof Plan", 1, page_break=True)
    add_body(doc, "The next milestone should not be a longer roadmap. It should be evidence that one community workflow is painful, repeatable, and worth paying for.")
    add_number(doc, "Choose one buyer and one workflow: for example, an RCCG province operations lead staffing a convention or managing verified volunteers across parishes.")
    add_number(doc, "Get a written pilot commitment with one named operational owner, one executive sponsor, a start date, and a simple success definition.")
    add_number(doc, "Onboard 100 to 300 real members or volunteers through the hierarchy. Track completion rate, drop-off points, and manual support time.")
    add_number(doc, "Issue a real credential and require it for one bounded event or opportunity. Measure scan success, exceptions, and operator workload.")
    add_number(doc, "Create at least one reuse event: a credential issued or earned in one context must affect an approval decision in another context.")
    add_number(doc, "Integrate one identity provider only where required. Compare provider cost, latency, failure rate, and manual-review burden.")
    add_number(doc, "Run an explicit willingness-to-pay conversation before the pilot ends. Ask for a paid renewal or rollout, not compliments.")
    add_table(
        doc,
        ["Metric", "Minimum evidence worth learning from", "What it answers"],
        [
            ["Pilot sponsor", "Named operator and budget-aware executive", "Is this a real organizational priority?"],
            ["Activated users", "100+ real members or volunteers", "Can onboarding work outside a demo?"],
            ["Credential reuse", "At least 20 cross-context verification decisions", "Is portability actually valuable?"],
            ["Operator time", "Measured baseline and post-pilot comparison", "Does the product remove work or add work?"],
            ["Paid continuation", "One signed renewal or rollout proposal", "Will anyone pay after novelty fades?"],
            ["Disputes", "Documented exception and appeals workflow", "Can the trust model survive real-world disagreement?"],
        ],
        [1.4, 2.6, 2.9],
        size=8.4,
    )

    add_heading(doc, "8. Product Scope Decisions", 1)
    add_table(
        doc,
        ["Decision", "Recommendation", "Reason"],
        [
            ["Estate billing, dues, and resident communications", "Do not build now", "Gate Africa, Sylogate, MyGate, and others already operate here."],
            ["Identity verification engine", "Integrate providers", "Prembly and Youverify already expose verification and screening capabilities."],
            ["Open artisan marketplace", "Do not build now", "It introduces a second business model and a cold-start problem."],
            ["Trust Passport and QR verification", "Build and test", "Central to the portability hypothesis."],
            ["Community hierarchy and authority delegation", "Build and test", "Potentially distinctive for RCCG and associations."],
            ["Opportunity publishing and eligibility gating", "Build and test", "Connects hierarchy and credentials to a real operational outcome."],
            ["Complex trust scoring", "Simplify", "Use explainable evidence before adding opaque scoring logic."],
            ["Procurement governance", "Defer", "Large surface area and mature adjacent products; revisit after wedge traction."],
        ],
        [2.2, 1.45, 3.25],
        size=8.35,
    )

    add_heading(doc, "9. Reference Sources", 1, page_break=True)
    add_body(doc, "Official product pages reviewed on 2 June 2026. Vendor-reported metrics and feature descriptions should be validated during customer discovery.")
    sources = [
        ("Gate Africa", "https://gate.africa/"),
        ("Gate Africa FAQ", "https://gate.africa/faq"),
        ("Sylogate", "https://www.sylogate.com/"),
        ("MyGate visitor management", "https://mygate.com/visitor-management/"),
        ("MyGate community management", "https://mygate.com/community-management/"),
        ("Rosterfy", "https://www.rosterfy.com/"),
        ("Rosterfy onboarding", "https://www.rosterfy.com/recruit-onboard/"),
        ("ComplyFlow", "https://www.complyflow.com/"),
        ("ComplyFlow workforce management", "https://www.complyflow.com/enterprise/workforce-management"),
        ("Appruv", "https://appruv.com/"),
        ("Checkr platform", "https://checkr.com/platform"),
        ("Prembly sandbox documentation", "https://docs.prembly.com/docs/environment"),
        ("Prembly background checks", "https://blog.prembly.com/prembly-background-check-feature/"),
        ("Youverify operating system", "https://youverify.co/product/youverify-operating-system"),
    ]
    add_table(doc, ["Source", "URL"], sources, [2.2, 4.8], size=8.35)

    add_heading(doc, "Internal Sources", 2)
    add_bullet(doc, "docs/20-vision-refactor.md: canonical June 2026 product positioning.")
    add_bullet(doc, "docs/07-mvp-scope.md: implemented MVP loop and post-MVP backlog.")
    add_bullet(doc, "docs/06-business-design.md: customer segments, pricing hypotheses, and revenue assumptions.")
    add_bullet(doc, "docs/17-competitive-analysis.md: prior internal competitor framing reviewed and challenged by this memo.")

    add_heading(doc, "Closing View", 1)
    add_body(doc, "TrustGrid should not try to win a slide-deck argument against every adjacent platform. It should prove one thing that existing products do not obviously prove: that an organized community will issue portable, authority-backed credentials and then use those credentials to make repeated decisions across its network.")
    add_body(doc, "If that behavior appears, the broader trust-graph story becomes credible. If it does not, the team should narrow or change the product before building the rest of the roadmap.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
