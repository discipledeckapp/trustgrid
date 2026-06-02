from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_SCRIPTS = Path(
    "/Users/seyiadelaju/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.430.10722/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights

OUT = Path("/Users/seyiadelaju/Projects/trustgrid/deliverables/trustgrid-seeded-demo-credentials.docx")

NAVY = "172554"
INDIGO = "EEF2FF"
AMBER = "FEF3C7"
GRAY = "F1F5F9"
MUTED = RGBColor(71, 85, 105)
TEXT = RGBColor(15, 23, 42)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_run(run, *, size=10, bold=False, color=TEXT):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_text(cell, text, *, bold=False, color=TEXT, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    style_run(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, weights, *, header_fill=GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9)
        set_cell_fill(table.rows[0].cells[idx], header_fill)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
    widths = column_widths_from_weights(weights)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    return table


def add_note(doc, title, text, *, fill=AMBER):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{title}: ")
    style_run(run, size=10, bold=True)
    run = paragraph.add_run(text)
    style_run(run, size=10)
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=0)
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(text)
    style_run(run, size=10)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, *, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        style_run(first, size=10, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        style_run(rest, size=10)
    else:
        run = paragraph.add_run(text)
        style_run(run, size=10)
    return paragraph


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(30, 41, 59)
    styles["Subtitle"].font.name = "Arial"
    styles["Subtitle"].font.size = Pt(11)
    styles["Subtitle"].font.color.rgb = MUTED
    for name, size in (("Heading 1", 15), ("Heading 2", 12)):
        styles[name].font.name = "Arial"
        styles[name].font.bold = True
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(30, 64, 175)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("TRUSTGRID  |  SEEDED DEMO ACCESS"), size=8, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("Confidential demo credentials  |  Rotate before any real launch"), size=8, color=MUTED)

    doc.add_heading("TrustGrid Seeded Demo Credentials", 0)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Local access guide and production deployment reference")
    meta = add_table(
        doc,
        ["Document scope", "Verified on", "Shared demo password"],
        [["Seeded personas and login URLs", "1 June 2026", "Admin123!"]],
        [2.4, 1.1, 1.2],
        header_fill=INDIGO,
    )
    doc.add_paragraph()

    add_note(
        doc,
        "Security warning",
        "This document contains demo credentials. Use it only for controlled demonstrations. "
        "Do not treat Admin123! as a production password. Rotate or disable all seeded accounts before a real launch.",
    )

    add_heading(doc, "1. Quick Start", 1)
    add_body(doc, "Use the phone number as the most portable login identifier. The seed also stores the email addresses listed below, and the login endpoint accepts either phone or email.")
    add_table(
        doc,
        ["Surface", "Local URL", "Configured production URL", "Primary persona"],
        [
            ["Super-admin console", "http://localhost:3003/login", "https://admin.trustgrid.ng/login", "Platform Admin"],
            ["Institution dashboard", "http://localhost:3001/login", "https://app.trustgrid.ng/login", "All tenant personas"],
            ["API documentation", "http://localhost:3000/api/docs", "https://api.trustgrid.ng/api/docs", "Developer reference"],
        ],
        [1.25, 2.05, 2.25, 1.35],
        header_fill=INDIGO,
    )

    add_heading(doc, "2. Complete Seeded Credential Matrix", 1)
    add_body(doc, "All personas below are created by backend/src/prisma/seed.ts. Every account uses the shared demo password Admin123!.")
    credential_rows = [
        ["PLATFORM_ADMIN", "Platform Admin", "08001234569", "admin@trustgrid.ng", "Super-admin console"],
        ["INSTITUTION_ADMIN", "Deacon Emeka", "08001234567", "emeka@redemptioncity.ng", "Institution dashboard"],
        ["INSTITUTION_OPERATOR", "Sister Adaeze", "08001234568", "adaeze@redemptioncity.ng", "Institution dashboard"],
        ["WORKER", "Chukwuemeka Adeyemi", "08001234570", "chukwuemeka@demo.com", "Institution dashboard / worker app"],
        ["WORKER", "Ngozi Okafor", "08001234572", "ngozi.okafor@legalaid.com", "Institution dashboard / worker app"],
        ["ORGANISATION_ADMIN", "Emeka Obi", "08001234573", "emeka.obi@emekaelectrical.ng", "Institution dashboard"],
        ["RESIDENT", "Funke Adeola", "08001234571", "funke@resident.com", "Institution dashboard"],
    ]
    add_table(
        doc,
        ["Role", "Persona", "Phone", "Seed email", "Login surface"],
        credential_rows,
        [1.35, 1.45, 1.15, 2.25, 1.35],
        header_fill=INDIGO,
    )

    add_heading(doc, "3. Local Verification Results", 1)
    add_body(doc, "The local API was tested directly on 1 June 2026. Each seeded persona authenticated successfully by phone with Admin123!. The local web login pages also returned HTTP 200.")
    add_table(
        doc,
        ["Check", "Result", "Notes"],
        [
            ["Super-admin page", "PASS", "http://localhost:3003/login returned HTTP 200"],
            ["Institution page", "PASS", "http://localhost:3001/login returned HTTP 200"],
            ["All seven phones", "PASS", "Each phone authenticated against http://localhost:3000/api/v1/auth/login"],
            ["Email login behavior", "PASS", "The API accepts email or phone; email matching is case-insensitive"],
        ],
        [1.7, 0.8, 4.8],
        header_fill=INDIGO,
    )
    add_note(
        doc,
        "Local database drift",
        "An existing local database may contain older demo email values from previous seeds. "
        "If an email differs locally, use the phone number or rerun the seed. The phones above were verified locally.",
        fill=GRAY,
    )

    heading = add_heading(doc, "4. Production Readiness", 1)
    heading.paragraph_format.page_break_before = True
    add_body(doc, "The codebase is configured for the production URLs below, but those DNS names did not resolve from this machine on 1 June 2026. Production sign-in could not be verified.")
    add_table(
        doc,
        ["Production endpoint", "Current verification", "Required action"],
        [
            ["https://admin.trustgrid.ng/login", "DNS unresolved", "Publish DNS and deploy the admin console"],
            ["https://app.trustgrid.ng/login", "DNS unresolved", "Publish DNS and deploy the tenant dashboard"],
            ["https://api.trustgrid.ng/api/v1", "DNS unresolved", "Publish DNS and deploy the API"],
        ],
        [2.7, 1.3, 3.3],
        header_fill=INDIGO,
    )

    add_heading(doc, "Production Deployment Checklist", 2)
    add_bullet(doc, "Deploy the API, tenant dashboard, and super-admin console.")
    add_bullet(doc, "Publish DNS records for api.trustgrid.ng, app.trustgrid.ng, and admin.trustgrid.ng.")
    add_bullet(doc, "Run the database migrations in the production environment.")
    add_bullet(doc, "Run the seed command only for a controlled demo environment: npm run db:seed.")
    add_bullet(doc, "Verify phone and email login against https://api.trustgrid.ng/api/v1/auth/login.")
    add_bullet(doc, "Rotate Admin123! or disable all seeded users before onboarding real users.")

    add_heading(doc, "5. Recommended Demo Flow", 1)
    add_table(
        doc,
        ["Step", "Persona", "Action"],
        [
            ["1", "Platform Admin", "Open the super-admin console and review platform statistics, institutions, workers, organisations, catalog, and passports."],
            ["2", "Deacon Emeka", "Open the institution dashboard and review the Redemption City workforce registry."],
            ["3", "Sister Adaeze", "Demonstrate day-to-day institution operations."],
            ["4", "Chukwuemeka Adeyemi", "Demonstrate an artisan worker account."],
            ["5", "Ngozi Okafor", "Demonstrate a professional worker profile."],
            ["6", "Emeka Obi", "Demonstrate organisation-admin access."],
            ["7", "Funke Adeola", "Demonstrate the resident/community-member persona."],
        ],
        [0.55, 1.75, 5.0],
        header_fill=INDIGO,
    )

    add_heading(doc, "6. Source of Truth", 1)
    add_body(doc, "Seed file: backend/src/prisma/seed.ts")
    add_body(doc, "Production URL configuration: admin/vercel.json, dashboard/vercel.json, and render.yaml")
    add_body(doc, "Authentication request shape: { identifier, password }, where identifier may be a phone number or email address.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
